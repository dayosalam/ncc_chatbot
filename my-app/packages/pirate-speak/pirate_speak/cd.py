# %%
import os
import docx


# def docx_to_txt(docx_path, txt_path):
#     """
#     Converts a .docx file to a .txt file.

#     Parameters:
#     docx_path (str): Path to the .docx file.
#     txt_path (str): Path to the output .txt file.
#     """
#     doc = docx.Document(docx_path)
#     with open(txt_path, 'w', encoding='utf-8') as txt_file:
#         for paragraph in doc.paragraphs:
#             txt_file.write(paragraph.text + '\n')


# def convert_all_docx_in_directory(source_directory, target_directory):
#     """
#     Converts all .docx files in a directory to .txt files and saves them in the target directory.

#     Parameters:
#     source_directory (str): Directory to search for .docx files.
#     target_directory (str): Directory to save the converted .txt files.
#     """
#     for root, _, files in os.walk(source_directory):
#         for file in files:
#             if file.endswith('.docx'):
#                 docx_path = os.path.join(root, file)
#                 txt_filename = os.path.splitext(file)[0] + '.txt'
#                 txt_path = os.path.join(target_directory, txt_filename)
#                 docx_to_txt(docx_path, txt_path)
#                 print(f"Converted {docx_path} to {txt_path}")


# # Create the target directory if it doesn't exist
# target_directory = r"C:\Users\i\Desktop\llm\ncc_doc\publish"
# os.makedirs(target_directory, exist_ok=True)

# # Source directory
# source_directory = r"C:\Users\i\Desktop\llm\ncc_doc"

# # Convert all .docx files in the source directory to .txt files in the target directory
# convert_all_docx_in_directory(source_directory, target_directory)


# %%

# Specify the directory containing the text files
# input_directory = r"C:\Users\i\Desktop\llm\ncc_doc\publish"

# # Specify the output file
# output_file = "combined_output2.txt"

# # Initialize a list to hold the contents of all files
# all_texts = []

# # Iterate over all files in the directory
# for filename in os.listdir(input_directory):
#     # Check if the file is a text file
#     if filename.endswith(".txt"):
#         file_path = os.path.join(input_directory, filename)
#         # Read the contents of the file
#         with open(file_path, 'r', encoding='utf-8') as file:
#             all_texts.append(file.read())

# # Combine all texts into a single string with a separator (optional)
# combined_text = "\n\n".join(all_texts)  # Using double newlines as a separator

# # Write the combined text to the output file
# with open(output_file, 'w', encoding='utf-8') as output:
#     output.write(combined_text)

# print(f"Combined text written to {output_file}")

# %%
from docx import Document

def txt_to_docx(txt_file_path, docx_file_path):
    # Create a new Document
    doc = Document()

    # Open and read the content of the text file
    with open(txt_file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    
    # Add each line to the Document
    for line in lines:
        doc.add_paragraph(line)

    # Save the Document as a .docx file
    doc.save(docx_file_path)
    print(f'Successfully converted {txt_file_path} to {docx_file_path}')

# Make sure to export the function
__all__ = ['txt_to_docx']
# Example usage
# txt_file = 'combined_output2.txt'
# docx_file = 'combined_output2.docx'
# txt_to_docx(txt_file, docx_file)

# import os
# import docx

# def docx_to_txt(docx_path, txt_path):
#     """
#     Converts a .docx file to a .txt file.

#     Parameters:
#     docx_path (str): Path to the .docx file.
#     txt_path (str): Path to the output .txt file.
#     """
#     doc = docx.Document(docx_path)
#     with open(txt_path, 'w', encoding='utf-8') as txt_file:
#         for paragraph in doc.paragraphs:
#             txt_file.write(paragraph.text + '\n')

# def convert_all_docx_in_directory(source_directory, target_directory):
#     """
#     Converts all .docx files in a directory to .txt files and saves them in the target directory.

#     Parameters:
#     source_directory (str): Directory to search for .docx files.
#     target_directory (str): Directory to save the converted .txt files.
#     """
#     for root, _, files in os.walk(source_directory):
#         for file in files:
#             if file.endswith('.docx'):
#                 docx_path = os.path.join(root, file)
#                 txt_filename = os.path.splitext(file)[0] + '.txt'
#                 txt_path = os.path.join(target_directory, txt_filename)
#                 docx_to_txt(docx_path, txt_path)
#                 print(f"Converted {docx_path} to {txt_path}")

# # Create the target directory if it doesn't exist
# target_directory = r"C:\Users\i\Desktop\llm\ncc_doc\publish\FAQs"
# os.makedirs(target_directory, exist_ok=True)

# # Source directory
# source_directory = r"C:\Users\i\Desktop\llm\ncc_doc\publish"

# Convert all .docx files in the source directory to .txt files in the target directory
# convert_all_docx_in_directory(source_directory, target_directory)