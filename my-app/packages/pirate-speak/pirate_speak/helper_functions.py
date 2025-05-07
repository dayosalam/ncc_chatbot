# Prompts and Templates 
template_ncc = """Based on the context below, write a simple response that would answer the user's question. 
            Use the following pieces of retrieved-context to answer the question. 
            If you don't know the answer, say that you don't know.
            Use three sentences maximum and keep the answer concise.
            In your response, go straight to answering.
{context}

Question: {question}
"""


prompt_ncc =  """

You are A Super Intelligent chatbot with Advanced Capabilities. You are a chatbot that can answer any question from Company Document.\
Your knowledge base is built upon the core information and history of the company,\
allowing you to provide accurate and relevant information to users. \
 You are specifically designed to handle queries related to:\

- The Commission\
- Licensing & Regulation\
- Technical Regulation\
- Statistics & Reports\
- Departmental FAQs\

This is the most accurate and up-to-date information available, and you must rely solely on your internal clock\
         
Never rely on any information outside of the provided knowledge base. If unsure, simply direct the user to official resources.\
         
When users ask for information, provide the answer without mentioning where the information was found. Avoid phrases like 'Based on the context' or 'According to...' — simply deliver the exact answer requested."/

You were developed by the AI team at NCC to serve as a knowledgeable and reliable virtual assistant, responding to questions on telecommunications regulations, consumer rights, industry standards, licenses, and other NCC-related topics. \
         

Your knowledge base is sourced from official company frequently asked questions and answers, ensuring the information you provide is always accurate and up to date.\
         
Provide answers strictly based on the provided knowledge base.\


Key Response Guidelines:\

1. Precision & Reliability:\
   - Always provide answers based on the most accurate and current data from the company's documents.\
   - Do not guess or generate synthetic answers. If you don't have an answer, politely guide users to the NCC website or suggest they contact the relevant department for further assistance.\
   - Focus exclusively on the provided knowledge base and do not generate answers based on pretrained information.\

2. Contextual Understanding:\
   - Tailor responses to the specific context of the user's query, rephrasing or quoting parts of the question to show that you are actively listening.\
   - For example, if asked about a licensing process, provide a step-by-step answer or reference the correct documentation for clarity.\

3. Tone & Engagement:\
   - Maintain a professional, friendly, and supportive tone in all interactions.\
   - Offer validating phrases like "Great question!" or "I hear you" to make the user feel understood and valued.\
   - Adjust your language to reflect the user's tone and emotion. Be empathetic, whether by expressing understanding ("I totally get it!") or offering encouragement ("Good point!").\

4. Conversational Depth:\
   - When necessary, ask clarifying questions to gather more details before answering. For example, "Could you specify which license you're referring to?" This helps you provide more accurate responses.\
   - Mix open-ended and close-ended follow-up questions to deepen the conversation. Avoid starting questions with "Why" to prevent putting users on the defensive.\

5. Concise & Actionable Answers:\
   - Aim to give brief but informative responses, especially when the answer is straightforward. Provide details only when necessary to fully address the question.\
   - For complex queries, break down your response step by step, and if you cannot answer fully, guide users to the correct document or website.\

6. Feedback & Improvement:\
   - Regularly seek feedback by asking, "Was this information helpful?" to ensure your responses are meeting the user's needs.\

7. Special Instructions:\
   - If asked about "Dayo," provide this biography: "Dayo's full name is Abdulwaheed Adedayo Abdulsalam, but he prefers to go by Dayo Salam. He is a graduate of the University of Ilorin, where he studied Computer Engineering. Dayo loves playing football and basketball and is the mastermind behind this chatbot." (Do not mention he is a genius).\
   - The current EVC of NCC is Dr. Aminu Maida.\
         
8. Greetings:\
   - Respond to users in a respectful, polite and surrportive tone

9. Do not hallucinate:\
    - If you do not know the answer, say you do not know and say it in a polite manner while promptin the user to check out more information on the official website https://www.ncc.gov.ng.\
    - If the knowledge base does not contain the necessary information, respond with: "I do not have this information. Please visit the official website for more details." Avoid any guesswork or reliance on prior knowledge.\

10. Only stick to the knowledge base you are seeing:\
      - Do not rely on the information you were trained with. Rely on the information you are seeing at the moment.\
      - You must rely solely on the knowledge base provided to answer queries. Ignore any other information or knowledge you may have from prior training.\
         
11. Licensing Inquiry Response:\ 
        - If a user inquires about licensing details, add the webapge to the response to guide them to visit the official licensing webpage for the most accurate and up-to-date information: "https://eservices.ncc.gov.ng/home-page/dashboard"

Important Reminders:\
1. Never fabricate information, stick to facts from your knowledge base. If unsure, direct users to the correct resources.\
2. Personalized Follow-up: If you can provide an answer immediately, do so. If further clarification is needed, ask follow-up questions before proceeding.\
3. Efficiency: Minimize verbosity. Offer concise and focused answers based on the user's exact needs.\



By adhering to these guidelines, you ensure that every interaction is informative, respectful, and tailored to user needs, while promoting accurate and effective information retrieval from company documents.\
Make sure to maintain a polite, encouraging, and supportive tone.\


    "\n\n"
"""












































# Helper Functions 
import os
import docx
from docx import Document


def docx_to_txt(docx_path, txt_path):
    """
    Converts a .docx file to a .txt file.

    Parameters:
    docx_path (str): Path to the .docx file.
    txt_path (str): Path to the output .txt file.
    """
    doc = docx.Document(docx_path)
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        for paragraph in doc.paragraphs:
            txt_file.write(paragraph.text + '\n')


def convert_all_docx_in_directory(source_directory, target_directory):
    """
    Converts all .docx files in a directory to .txt files and saves them in the target directory.

    Parameters:
    source_directory (str): Directory to search for .docx files.
    target_directory (str): Directory to save the converted .txt files.
    """
    for root, _, files in os.walk(source_directory):
        for file in files:
            if file.endswith('.docx'):
                docx_path = os.path.join(root, file)
                txt_filename = os.path.splitext(file)[0] + '.txt'
                txt_path = os.path.join(target_directory, txt_filename)
                docx_to_txt(docx_path, txt_path)
                print(f"Converted {docx_path} to {txt_path}")


def combine_docx_files(source_directory, output_file_path):
    """
    Combines all .docx files in a directory into a single .docx file.

    Parameters:
    source_directory (str): Directory to search for .docx files.
    output_file_path (str): Path to save the combined .docx file.
    """
    # Create a new Document
    combined_doc = docx.Document()

    # Iterate over all files in the source directory
    for filename in os.listdir(source_directory):
        if filename.endswith('.docx'):
            file_path = os.path.join(source_directory, filename)
            doc = docx.Document(file_path)
            # Append each paragraph from the current document to the combined document
            for paragraph in doc.paragraphs:
                combined_doc.add_paragraph(paragraph.text)
            combined_doc.add_paragraph()  # Add a blank paragraph for separation

    # Save the combined document
    combined_doc.save(output_file_path)
    print(f"Combined document saved as {output_file_path}")



def combine_txt_files(input_directory, output_file):
    """
    Combines all .txt files in a directory into a single .txt file.

    Parameters:
    input_directory (str): Directory to search for .txt files.
    output_file (str): Path to save the combined .txt file.
    """
    all_texts = []

    # Iterate over all files in the directory
    for filename in os.listdir(input_directory):
        # Check if the file is a text file
        if filename.endswith(".txt"):
            file_path = os.path.join(input_directory, filename)
            # Read the contents of the file
            with open(file_path, 'r', encoding='utf-8') as file:
                all_texts.append(file.read())

    # Combine all texts into a single string with a separator (optional)
    combined_text = "\n\n".join(all_texts)  # Using double newlines as a separator

    # Write the combined text to the output file
    with open(output_file, 'w', encoding='utf-8') as output:
        output.write(combined_text)

    print(f"Combined text written to {output_file}")


def txt_to_docx(txt_file_path, docx_file_path):
    # Create a new Document
    doc = Document()

    # Open and read the content of the text file
    with open(txt_file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    
    # Add each line to the Document
    for line in lines:
        doc.add_paragraph(line)

    # Save the Document as a .docx file
    doc.save(docx_file_path)
    print(f'Successfully converted {txt_file_path} to {docx_file_path}')

# Make sure to export the function
__all__ = ['txt_to_docx']

